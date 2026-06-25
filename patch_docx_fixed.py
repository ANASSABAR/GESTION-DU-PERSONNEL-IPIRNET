import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document(r'C:\xampp\htdocs\IP\Dictionnaire_Donnees_Gestion_Personnel_Final.docx')

def insert_paragraph_after(element, text, bold=False):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    r.append(t)
    p.append(r)
    element.addnext(p)

def add_row(table, data):
    row = table.add_row()
    for i, val in enumerate(data):
        row.cells[i].text = val

# 1. HEURE_SUPPLEMENTAIRE (Table 4)
t_hs = doc.tables[4]
data_hs = [
    ['id_heure_sup', 'INT', '11', 'PK', 'NOT NULL'],
    ['date', 'DATE', '-', '', 'NOT NULL'],
    ['nombre_heures', 'DECIMAL', '5,2', '', 'NOT NULL'],
    ['contrat', 'VARCHAR', '50', '', 'NULL'],
    ['prix_heure', 'DECIMAL', '10,2', '', 'NULL'],
    ['motif', 'VARCHAR', '255', '', 'NULL'],
    ['id_personnel', 'INT', '11', 'FK', 'NOT NULL']
]
for row in t_hs.rows[1:]:
    t_hs._tbl.remove(row._tr)
for d in data_hs:
    add_row(t_hs, d)

insert_paragraph_after(t_hs._tbl, "Règles métier : CDI = 25 DH/h, CDD = 20 DH/h, STAGE = interdit heures supplémentaires, INTERIM = interdit heures supplémentaires")

# 2. PERSONNEL (Table 5)
t_pers = doc.tables[5]
data_pers = [
    ['id_personnel', 'INT', '11', 'PK', 'NOT NULL'],
    ['nom', 'VARCHAR', '100', '', 'NOT NULL'],
    ['prenom', 'VARCHAR', '100', '', 'NOT NULL'],
    ['telephone', 'VARCHAR', '20', '', 'NULL'],
    ['email', 'VARCHAR', '150', 'UNIQUE', 'NULL'],
    ['adresse', 'TEXT', '-', '', 'NULL'],
    ['cin', 'VARCHAR', '20', '', 'NULL'],
    ['date_naissance', 'DATE', '-', '', 'NULL'],
    ['sexe', 'VARCHAR', '10', '', 'NULL'],
    ['contrat', 'VARCHAR', '50', '', 'NULL'],
    ['date_recrutement', 'DATE', '-', '', 'NULL'],
    ['statut', 'ENUM', "'Actif','Essai','Congé','Inactif'", '', 'NOT NULL'],
    ['id_categorie', 'INT', '11', 'FK', 'NOT NULL']
]
for row in t_pers.rows[1:]:
    t_pers._tbl.remove(row._tr)
for d in data_pers:
    add_row(t_pers, d)

# 3. REMUNERATION (Table 6)
t_rem = doc.tables[6]
data_rem = [
    ['id_remuneration', 'INT', '11', 'PK', 'NOT NULL'],
    ['salaire_base', 'DECIMAL', '10,2', '', 'NOT NULL'],
    ['prime', 'DECIMAL', '10,2', '', 'NOT NULL'],
    ['quantite_heures_supp', 'DECIMAL', '5,2', '', 'NULL'],
    ['prix_unitaire_heure', 'DECIMAL', '10,2', '', 'NULL'],
    ['montant_total_heures_supp', 'DECIMAL', '10,2', '', 'NULL'],
    ['deduction_cnss', 'DECIMAL', '10,2', '', 'NOT NULL'],
    ['salaire_net', 'DECIMAL', '10,2', '', 'NULL'],
    ['date_paiement', 'DATE', '-', '', 'NULL'],
    ['id_personnel', 'INT', '11', 'FK', 'NOT NULL']
]
for row in t_rem.rows[1:]:
    t_rem._tbl.remove(row._tr)
for d in data_rem:
    add_row(t_rem, d)

insert_paragraph_after(t_rem._tbl, "Formule : Salaire Net = Salaire Base + Prime + Montant HS - Déduction CNSS")

# 4. Ajouter section finale "Contraintes métier"
# Using safe formatting instead of named styles
p_heading = doc.add_paragraph()
r_heading = p_heading.add_run('Contraintes métier')
r_heading.bold = True

doc.add_paragraph("- Contrat ∈ {CDI, CDD, STAGE, INTERIM}")
doc.add_paragraph("- Motif ∈ {Réunion pédagogique, Surveillance d’examen, Révision programme, Correction copies, Formation interne, Maintenance système, Support administratif, Travail exceptionnel, Urgence service, Projet spécial}")

doc.save(r'C:\xampp\htdocs\IP\Dictionnaire_Donnees_Gestion_Personnel_Final.docx')
print('Docx saved successfully.')
