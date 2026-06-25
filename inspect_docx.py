import docx

doc = docx.Document(r'C:\xampp\htdocs\IP\Dictionnaire_Donnees_Gestion_Personnel_Final.docx')

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print("\n=== TABLES ===")
for i, t in enumerate(doc.tables):
    print(f"Table {i}:")
    for j, row in enumerate(t.rows):
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"  Row {j}: {cells}")
